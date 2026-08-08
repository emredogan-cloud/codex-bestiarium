#!/usr/bin/env python3
"""
CODEX BESTIARIUM — KINDLE EPUB (reflowable) VE DOCX YEDEĞİ
================================================================================
    NEDEN YENİ BİR BETİK
    ────────────────────
    `make_docx_epub.py` Cilt 1'den devralındı ve `CODEX_MYTHOLOGICA.epub`
    üretir; yapısı da o kitabındır (uygarlık → mit). Bestiarium'un yapısı
    sınıf → madde ve her maddede yedi bölüm. Aynı gerekçeyle `make_book.py`
    yazıldıysa burada da yenisi yazılır.

REFLOWABLE — SABİT DÜZEN DEĞİL
    Yol haritası bunu iki kez söylüyor ve gerekçesi finansal: sabit düzen
    Kindle Translate kapısını KALICI olarak kapatır. Bu dosyada sabit
    boyut, sabit konum ve mutlak birim yoktur; ölçüler `em` ve `%`.

PLAKA — İKİ TON, 900 px (D49)
    `plates_kindle` zaten bu iş için üretiliyor: iki tonlu PNG, plaka başına
    ~39 KB. Gravür dilinde tonu gri seviyeler değil TARAMANIN KENDİSİ taşır
    ve basılı gravür de iki tonludur. Bütçe: EPUB ≤ 7 MB.

ÇIKTILAR
    05_KINDLE/CODEX_BESTIARIUM.epub
    02_MANUSCRIPT/CODEX_BESTIARIUM_MANUSCRIPT.docx      (--docx)

ÇIKIŞ KODLARI
    0  üretildi      1  bütçe aşıldı veya kusur      2  metin yok

KULLANIM
    python3 08_BUILD/make_ebook.py
    python3 08_BUILD/make_ebook.py --docx
"""

from __future__ import annotations

import argparse
import html
import json
import os
import sys
import uuid
import zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import front_matter as FM  # noqa: E402
from bestiarium import (  # noqa: E402
    AUTHOR,
    BOOK_SUBTITLE,
    BOOK_TITLE,
    CLASS_IDS,
    ENTRY_SECTIONS,
    EPUB_BUDGET_MB,
    IMPRINT,
    KIN_IDS,
    MATTER_SECTIONS,
    ROOT,
    SERIES,
    VOLUME,
    load_book,
    load_spec,
    matter_group,
    region_en,
)

KINDLE_DIR = os.path.join(ROOT, "05_KINDLE")
PLATE_DIR = os.path.join(ROOT, "07_ASSETS", "plates_kindle")
EPUB_PATH = os.path.join(KINDLE_DIR, "CODEX_BESTIARIUM.epub")
DOCX_PATH = os.path.join(ROOT, "02_MANUSCRIPT",
                         "CODEX_BESTIARIUM_MANUSCRIPT.docx")

CSS = """
html { font-size: 100%; }
body { margin: 0 5%; line-height: 1.5; text-align: justify;
       widows: 2; orphans: 2; }
h1 { font-size: 1.6em; text-align: center; margin: 2em 0 0.2em;
     font-weight: normal; letter-spacing: 0.08em; }
h2 { font-size: 1.25em; text-align: center; margin: 1.6em 0 0.3em;
     font-weight: normal; letter-spacing: 0.06em; }
h3 { font-size: 1em; margin: 1.4em 0 0.3em; font-weight: bold; }
p  { margin: 0; text-indent: 1.2em; }
p.first, p.noindent { text-indent: 0; margin-top: 0.6em; }
p.opening { text-indent: 0; font-style: italic; margin: 0.8em 0; }
p.meta { text-indent: 0; text-align: center; font-size: 0.82em;
         color: #555; margin-bottom: 1.2em; }
p.kin { text-indent: 0; font-size: 0.9em; margin-top: 1em;
        border-top: 1px solid #8A6E2F; padding-top: 0.6em; }
p.sources { text-indent: 0; font-size: 0.85em; font-style: italic;
            color: #444; margin-top: 0.8em; }
p.note { text-indent: 0; font-style: italic; color: #555;
         margin-bottom: 1em; }
p.copy { text-indent: 0; font-size: 0.85em; margin-bottom: 0.7em; }
p.center { text-indent: 0; text-align: center; }
div.plate { text-align: center; margin: 0 0 1em; page-break-inside: avoid; }
div.plate img { max-width: 78%; height: auto; }
ul.idx { list-style: none; padding-left: 0; }
ul.idx li { text-indent: 0; margin-bottom: 0.35em; font-size: 0.92em;
            text-align: left; }
"""


def esc(t: str) -> str:
    return html.escape(t or "", quote=False)


def paras(text: str, cls_first: str = "first") -> str:
    out = []
    for i, p in enumerate(x.strip() for x in (text or "").split("\n\n")):
        if not p:
            continue
        if p.startswith("## "):
            out.append(f"<h3>{esc(p[3:].strip())}</h3>")
            continue
        cls = f' class="{cls_first}"' if i == 0 else ""
        out.append(f"<p{cls}>{esc(p)}</p>")
    return "\n".join(out)


def page(title: str, body: str) -> bytes:
    return (f'<?xml version="1.0" encoding="utf-8"?>\n'
            f'<!DOCTYPE html>\n'
            f'<html xmlns="http://www.w3.org/1999/xhtml" '
            f'xmlns:epub="http://www.idpf.org/2007/ops" lang="en">\n'
            f'<head><meta charset="utf-8"/><title>{esc(title)}</title>'
            f'<link rel="stylesheet" type="text/css" href="style.css"/>'
            f'</head>\n<body>\n{body}\n</body>\n</html>\n').encode("utf-8")


def build_epub(verbose: bool = True) -> int:
    book = load_book()
    if book is None or not book.get("entries"):
        print("ATLANDI: metin yok.")
        return 2
    spec = load_spec()
    os.makedirs(KINDLE_DIR, exist_ok=True)

    trads = {t["id"]: t for t in spec["traditions"]}
    classes = {c["id"]: c for c in spec["classes"]}
    kinfam = {f["id"]: f for f in spec["kinFamilies"]}
    by_id = {c["id"]: c for c in spec["creatures"]}
    creatures = sorted(spec["creatures"], key=lambda c: c.get("number", 0))

    files: list[tuple[str, bytes]] = []   # (ad, içerik)
    spine: list[tuple[str, str]] = []     # (dosya, başlık)
    images: list[str] = []

    def add(name: str, title: str, body: str):
        files.append((name, page(title, body)))
        spine.append((name, title))

    # ── başlık ve künye ─────────────────────────────────────────────────
    tp = FM.TITLE_PAGE
    add("title.xhtml", BOOK_TITLE,
        f'<h1>{esc(tp["title"])}</h1>'
        f'<p class="center"><em>{esc(tp["subtitle"])}</em></p>'
        f'<p class="center">{esc(tp["line"])}</p>'
        f'<p class="center" style="margin-top:2em">{esc(tp["author"])}</p>'
        f'<p class="center">{esc(tp["series"])}</p>'
        f'<p class="center">{esc(tp["imprint"])}</p>')
    add("copyright.xhtml", "Copyright",
        "\n".join(f'<p class="copy">{esc(x)}</p>' for x in FM.COPYRIGHT))
    add("dedication.xhtml", "Dedication",
        f'<p class="center"><em>{esc(FM.DEDICATION)}</em></p>')

    # ── kırk gelenek ────────────────────────────────────────────────────
    rows = []
    region = None
    for t in sorted(spec["traditions"],
                    key=lambda x: (region_en(x["regionGroup"]), x["name"])):
        if t["regionGroup"] != region:
            region = t["regionGroup"]
            rows.append(f"<h3>{esc(region_en(region))}</h3>")
        n = sum(1 for c in creatures if c["tradition"] == t["id"])
        rows.append(f'<p class="noindent">{esc(t["name"])} — {n}</p>')
    add("traditions.xhtml", FM.MAP_TITLE,
        f"<h2>{esc(FM.MAP_TITLE)}</h2>"
        f'<p class="note">{esc(FM.MAP_NOTE)}</p>' + "\n".join(rows))

    # ── giriş ve nasıl okunur ───────────────────────────────────────────
    for group, key, title, _p in MATTER_SECTIONS:
        if group != "front":
            continue
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if item.get("body"):
            add(f"m-{key}.xhtml", title,
                f"<h2>{esc(title)}</h2>" + paras(item["body"]))

    # ── sınıflar ────────────────────────────────────────────────────────
    kin_home: dict[str, list[str]] = {}
    for fid in KIN_IDS:
        mem = [c for c in creatures if c.get("kinFamily") == fid]
        if not mem:
            continue
        cnt: dict[str, int] = {}
        for m in mem:
            cnt[m["class"]] = cnt.get(m["class"], 0) + 1
        kin_home.setdefault(max(cnt, key=cnt.get), []).append(fid)

    for cid in CLASS_IDS:
        kl = classes[cid]
        members = [c for c in creatures if c["class"] == cid]
        add(f"class-{cid}.xhtml", f"{cid} · {kl['en']}",
            f"<h1>{esc(kl['en'])}</h1>"
            f'<p class="meta">Class {cid} · {len(members)} creatures</p>'
            + paras((book.get("classOpenings") or {}).get(cid, "")))

        for fid in kin_home.get(cid, []):
            fam = kinfam[fid]
            mem = [c for c in creatures if c.get("kinFamily") == fid]
            add(f"kin-{fid}.xhtml", fam["en"],
                f"<h1>{esc(fam['en'])}</h1>"
                f'<p class="meta">Kin image {fid} · {len(mem)} creatures · '
                f'{len({m["tradition"] for m in mem})} traditions · '
                f'{esc(fam.get("motif",""))}</p>'
                + paras((book.get("kinOpenings") or {}).get(fid, "")))

        for rec in members:
            entry = (book.get("entries") or {}).get(rec["id"])
            if not entry:
                continue
            sec = entry["sections"]
            plate = os.path.join(PLATE_DIR, rec.get("plate", "") + ".png")
            body = []
            if os.path.exists(plate):
                images.append(plate)
                body.append(f'<div class="plate"><img src="images/'
                            f'{os.path.basename(plate)}" alt=""/></div>')
            body.append(f"<h2>{esc(rec['name'])}</h2>")
            fam = kinfam.get(rec.get("kinFamily"))
            meta = (f"{trads.get(rec['tradition'],{}).get('name','')} · "
                    f"{rec['class']} · {kl['en']}")
            if fam:
                meta += f" · {fam['id']} · {fam['en']}"
            meta += f" · {' '.join(rec['motif'])}"
            if rec.get("pronunciation"):
                meta += f" · [{rec['pronunciation']}]"
            body.append(f'<p class="meta">{esc(meta)}</p>')
            for key, _l, _lo, _hi in ENTRY_SECTIONS:
                t = sec.get(key, "")
                if not t:
                    continue
                if key == "opening":
                    body.append(f'<p class="opening">{esc(t)}</p>')
                elif key == "kin":
                    lead = f"<b>{esc(fam['en'])}</b> — " if fam else ""
                    names = ", ".join(esc(by_id[r]["name"])
                                      for r in rec.get("crossRefs", [])
                                      if r in by_id)
                    body.append(f'<p class="kin">{lead}{esc(t)} '
                                f"<i>See also:</i> {names}.</p>")
                elif key == "sources":
                    body.append(f'<p class="sources">{esc(t)}</p>')
                else:
                    cls = ' class="first"' if key == "where" else ""
                    body.append(f"<p{cls}>{esc(t)}</p>")
            add(f"e-{rec['id']}.xhtml", rec["name"], "\n".join(body))

    # ── sonsöz ve arka madde ────────────────────────────────────────────
    for group, key, title, _p in MATTER_SECTIONS:
        if group != "back":
            continue
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if item.get("body"):
            add(f"m-{key}.xhtml", title,
                f"<h2>{esc(title)}</h2>" + paras(item["body"]))

    # ── kaynaklar ───────────────────────────────────────────────────────
    src = [f"<h2>{esc(FM.SOURCES_TITLE)}</h2>",
           f'<p class="note">{esc(FM.SOURCES_NOTE)}</p>']
    for t in sorted(spec["traditions"], key=lambda x: x["name"]):
        mem = [c for c in creatures if c["tradition"] == t["id"]]
        if not mem:
            continue
        src.append(f"<h3>{esc(t['name'])}</h3>")
        for rec in mem:
            e = (book.get("entries") or {}).get(rec["id"])
            note = (e.get("sections") or {}).get("sources", "") if e else ""
            if note:
                src.append(f'<p class="noindent"><b>{esc(rec["name"])}</b> — '
                           f"{esc(note)}</p>")
    add("sources.xhtml", FM.SOURCES_TITLE, "\n".join(src))

    # ── DİZİN YOK, ARAMA VAR ────────────────────────────────────────────
    # Basılı kitabın dört dizini SAYFA NUMARASINA dayanır. Reflowable bir
    # e-kitapta sayfa numarası yoktur; okur cihazına göre değişir. Sayfa
    # numarasız bir dizin basmak, okura işe yaramaz bir liste vermektir.
    # Yerine geleneğe göre gruplanmış İÇİNDEKİLER ve cihazın kendi araması
    # kullanılır — Kindle'da doğru olan budur.

    uid = f"urn:uuid:{uuid.uuid5(uuid.NAMESPACE_URL, EPUB_PATH)}"
    nav = ['<nav epub:type="toc" id="toc"><h1>Contents</h1><ol>']
    for name, title in spine:
        nav.append(f'<li><a href="{name}">{esc(title)}</a></li>')
    nav.append("</ol></nav>")
    files.append(("nav.xhtml", page("Contents", "\n".join(nav))))

    manifest = ['<item id="nav" href="nav.xhtml" '
                'media-type="application/xhtml+xml" properties="nav"/>',
                '<item id="css" href="style.css" media-type="text/css"/>']
    for i, (name, _t) in enumerate(spine):
        manifest.append(f'<item id="x{i}" href="{name}" '
                        f'media-type="application/xhtml+xml"/>')
    seen = []
    for p in images:
        b = os.path.basename(p)
        if b in seen:
            continue
        seen.append(b)
        manifest.append(f'<item id="img{len(seen)}" href="images/{b}" '
                        f'media-type="image/png"/>')
    itemrefs = "".join(f'<itemref idref="x{i}"/>' for i in range(len(spine)))

    opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0"
         unique-identifier="bookid" xml:lang="en">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="bookid">{uid}</dc:identifier>
    <dc:title>{esc(BOOK_TITLE)}: {esc(BOOK_SUBTITLE)}</dc:title>
    <dc:creator>{esc(AUTHOR)}</dc:creator>
    <dc:publisher>{esc(IMPRINT)}</dc:publisher>
    <dc:language>en</dc:language>
    <dc:description>{esc(SERIES)} volume {VOLUME}</dc:description>
    <meta property="dcterms:modified">2026-08-08T00:00:00Z</meta>
    <meta property="belongs-to-collection" id="col">{esc(SERIES)}</meta>
    <meta refines="#col" property="collection-type">series</meta>
    <meta refines="#col" property="group-position">{VOLUME}</meta>
  </metadata>
  <manifest>{''.join(manifest)}</manifest>
  <spine>{itemrefs}</spine>
</package>
"""
    container = ('<?xml version="1.0"?>\n<container version="1.0" '
                 'xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
                 '<rootfiles><rootfile full-path="OEBPS/content.opf" '
                 'media-type="application/oebps-package+xml"/></rootfiles>'
                 "</container>")

    with zipfile.ZipFile(EPUB_PATH, "w") as z:
        z.writestr("mimetype", "application/epub+zip",
                   compress_type=zipfile.ZIP_STORED)
        z.writestr("META-INF/container.xml", container,
                   compress_type=zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/content.opf", opf,
                   compress_type=zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/style.css", CSS,
                   compress_type=zipfile.ZIP_DEFLATED)
        for name, data in files:
            z.writestr(f"OEBPS/{name}", data,
                       compress_type=zipfile.ZIP_DEFLATED)
        for b in seen:
            z.write(os.path.join(PLATE_DIR, b), f"OEBPS/images/{b}",
                    compress_type=zipfile.ZIP_DEFLATED)

    mb = os.path.getsize(EPUB_PATH) / 1048576
    if verbose:
        print("=" * 78)
        print("KINDLE EPUB (reflowable)")
        print("=" * 78)
        print(f"  bölüm dosyası : {len(spine)}")
        print(f"  plaka         : {len(seen)}")
        print(f"  boyut         : {mb:.2f} MB   (bütçe {EPUB_BUDGET_MB} MB)")
        print(f"\n  yazıldı: {os.path.relpath(EPUB_PATH, ROOT)}")
    if mb > EPUB_BUDGET_MB:
        print(f"\n[FAIL] EPUB bütçeyi aştı: {mb:.2f} > {EPUB_BUDGET_MB} MB")
        return 1
    if verbose:
        print(f"\n[  ok ] bütçe içinde · reflowable · sabit düzen yok")
    return 0


def build_docx() -> int:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("ATLANDI: python-docx yok — DOCX üretilmedi.")
        return 2
    book = load_book()
    spec = load_spec()
    if book is None:
        return 2
    creatures = sorted(spec["creatures"], key=lambda c: c.get("number", 0))
    classes = {c["id"]: c for c in spec["classes"]}

    doc = Document()
    cp = doc.core_properties
    cp.title = f"{BOOK_TITLE}: {BOOK_SUBTITLE}"
    cp.author = AUTHOR
    cp.language = "en-GB"
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(11)

    doc.add_heading(BOOK_TITLE, 0)
    doc.add_paragraph(BOOK_SUBTITLE)
    for line in FM.COPYRIGHT:
        doc.add_paragraph(line)
    for group, key, title, _p in MATTER_SECTIONS:
        if group != "front":
            continue
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if not item.get("body"):
            continue
        doc.add_page_break()
        doc.add_heading(title, 1)
        for p in item["body"].split("\n\n"):
            if p.strip():
                doc.add_paragraph(p.strip().removeprefix("## "))
    for cid in CLASS_IDS:
        doc.add_page_break()
        doc.add_heading(f"{cid} · {classes[cid]['en']}", 1)
        for p in (book.get("classOpenings") or {}).get(cid, "").split("\n\n"):
            if p.strip():
                doc.add_paragraph(p.strip())
        for rec in [c for c in creatures if c["class"] == cid]:
            e = (book.get("entries") or {}).get(rec["id"])
            if not e:
                continue
            doc.add_page_break()
            doc.add_heading(rec["name"], 2)
            for key, _l, _lo, _hi in ENTRY_SECTIONS:
                t = e["sections"].get(key, "")
                if t:
                    doc.add_paragraph(t)
    for group, key, title, _p in MATTER_SECTIONS:
        if group != "back":
            continue
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if not item.get("body"):
            continue
        doc.add_page_break()
        doc.add_heading(title, 1)
        for p in item["body"].split("\n\n"):
            if p.strip():
                doc.add_paragraph(p.strip().removeprefix("## "))
    os.makedirs(os.path.dirname(DOCX_PATH), exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"yazıldı: {os.path.relpath(DOCX_PATH, ROOT)} "
          f"({os.path.getsize(DOCX_PATH)/1048576:.2f} MB)")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--docx", action="store_true")
    args = ap.parse_args()
    rc = build_epub()
    if args.docx:
        # DOCX bir YEDEKTİR (yol haritası çıktı 6): python-docx yoksa
        # atlanır ve EPUB'ın sonucunu bozmaz.
        d = build_docx()
        if d == 1:
            rc = 1
    return rc


if __name__ == "__main__":
    raise SystemExit(main())
