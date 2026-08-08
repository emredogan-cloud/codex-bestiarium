#!/usr/bin/env python3
"""
CODEX BESTIARIUM — ANA DİL EDİTÖRÜ TESLİM PAKETİ (Faz 5 · Geçiş 3)
================================================================================
Yol haritası Geçiş 3'ü şöyle tanımlıyor: *"Ana dili İngilizce bir satır
editörü — ses doğallığı. Dışarıya verilecek İLK iştir."* Kurucunun Faz 5
emri de aynı yerde duruyor: insan editör tek izinli dış bağımlılıktır ve
**beklenmez** — metin hazırlanır, teslim belgesi yazılır, insan gözü
gereken bölümler işaretlenir.

Bu betik o paketi ÜRETİR. Elle hazırlanmaz, çünkü elle hazırlanan bir
paket bir sonraki turda bayatlar ve kimse fark etmez.

PAKET
    02_MANUSCRIPT/codex-bestiarium-editor.docx   düzenlenebilir metin
    02_MANUSCRIPT/codex-bestiarium-editor.md     DOCX yoksa veya yanında
    02_MANUSCRIPT/EDITOR_BRIEF.md                editörün brifingi
    01_SOURCE/editor_pack.json                   DEPODA kalan ölçü

    İlk üçü PROZA taşır ve `.gitignore`'dadır (karar A1/D29). Depoda
    kalan yalnızca sayılardır.

İNSAN GÖZÜ GEREKEN BÖLÜMLER — TÜRETİLİR, YAZILMAZ
    ① Defterin dokunduğu bölümler. Yeniden kurulmuş cümle, ana dili
      İngilizce olmayan bir yazarın en çok tökezlediği yerdir.
    ② Yaşayan gelenek kısıt cümleleri. Hem doğru hem doğal olmak
      zorundalar ve ikisi çatışırsa doğruluk kazanır — editör bunu
      bilmeden düzeltirse etik kapı bozulur.
    ③ Bant kenarındaki bölümler (%4 içinde). Sıkıştırma izi taşırlar.
    ④ Ritim aykırıları: blok cümle ortalaması dağılımın en alt veya en
      üst %5'inde (ölçülen eşik: 12,3 ve 17,8).
    ⑤ Ön ve arka madde. Yeni proza ve kitabın en çok okunan sayfaları.
    ⑥ Diakritik yoğun bölümler — dizgide en kolay bozulan yer.

KULLANIM
    python3 08_BUILD/editor_pack.py            # md + brief + ölçü
    python3 08_BUILD/editor_pack.py --docx     # ayrıca DOCX (python-docx)
    python3 08_BUILD/editor_pack.py --check    # ölçü güncel mi
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import unicodedata

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from bestiarium import (  # noqa: E402
    AUTHOR,
    BOOK_SUBTITLE,
    BOOK_TITLE,
    EDITOR_COPY_STEM,
    ENTRY_SECTIONS,
    FORBIDDEN_GAME_TERMS,
    FORBIDDEN_PHRASES,
    FORBIDDEN_SOFTENERS,
    FORBIDDEN_SUPERLATIVES,
    MATTER_SECTIONS,
    ROOT,
    WORD_BAND,
    load_book,
    load_spec,
    matter_group,
)
from textutil import sentences, word_count  # noqa: E402

MANUSCRIPT_DIR = os.path.join(ROOT, "02_MANUSCRIPT")
MD_PATH = os.path.join(MANUSCRIPT_DIR, f"{EDITOR_COPY_STEM}.md")
DOCX_PATH = os.path.join(MANUSCRIPT_DIR, f"{EDITOR_COPY_STEM}.docx")
BRIEF_PATH = os.path.join(MANUSCRIPT_DIR, "EDITOR_BRIEF.md")
MEASURE_PATH = os.path.join(ROOT, "01_SOURCE", "editor_pack.json")

# EŞİKLER ÖLÇÜLEREK SEÇİLDİ, tahmin edilmedi. İlk deneme 784 bölümün
# 363'ünü işaretledi — %46'lık bir "öncelik listesi" öncelik listesi
# değildir, kopyanın kendisidir ve editörün zamanını hiçbir yere
# yöneltmez.
#
# Blok cümle ortalamasının dağılımı ölçüldü (560 blok, ortalama 14,9):
#   %5 → 12,3   %10 → 12,6   %90 → 17,2   %95 → 17,8
# Yani 13–19 aralığı dağılımın ORTASINI işaretliyordu. Eşik uçlara
# çekildi: yaklaşık en alt ve en üst %5.
BAND_EDGE = 0.04          # bandın %4'ü kadar kenar payı
SENT_LOW, SENT_HIGH = 12.3, 17.8
DIACRITIC_RATE = 0.020    # bloktaki diakritikli harf oranı eşiği

SECTION_LABEL = {k: label for k, label, _, _ in ENTRY_SECTIONS}
SECTION_BANDS = {k: (lo, hi) for k, _, lo, hi in ENTRY_SECTIONS}


def diacritic_rate(text: str) -> float:
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return 0.0
    marked = sum(1 for c in letters
                 if len(unicodedata.normalize("NFD", c)) > 1)
    return marked / len(letters)


def flag_sections(book: dict, spec: dict, touched: set[str]) -> list[dict]:
    """İnsan gözü gereken bölümler — hepsi ölçümden türetilir."""
    flags: list[dict] = []

    def add(ref, reason, detail=""):
        flags.append({"ref": ref, "reason": reason, "detail": detail})

    restricted = {c["id"] for c in spec["creatures"]
                  if c.get("restrictionScreened")}

    for cid, entry in (book.get("entries") or {}).items():
        for key, body in (entry.get("sections") or {}).items():
            if not body:
                continue
            ref = f"{cid}/{key}"
            if ref in touched:
                add(ref, "defter dokundu",
                    "cümle yeniden kuruldu; doğallık denetimi gerekiyor")
            lo, hi = SECTION_BANDS[key]
            n = word_count(body)
            pad = max(1, round((hi - lo) * BAND_EDGE))
            if n <= lo + pad or n >= hi - pad:
                add(ref, "bant kenarı", f"{n} kelime · bant {lo}–{hi}")
            s = sentences(body)
            if s and key != "opening":
                avg = sum(word_count(x) for x in s) / len(s)
                if avg < SENT_LOW or avg > SENT_HIGH:
                    add(ref, "ritim aykırısı",
                        f"ortalama cümle {avg:.1f} kelime")
            if diacritic_rate(body) >= DIACRITIC_RATE:
                add(ref, "diakritik yoğun",
                    f"{diacritic_rate(body) * 100:.1f}% işaretli harf")
        if cid in restricted:
            body = (entry.get("sections") or {}).get("does", "")
            where = (entry.get("sections") or {}).get("where", "")
            for key, text in (("where", where), ("does", body)):
                if not text:
                    continue
                # DAR TUTULUYOR. İlk deneme "published" ve "belong"
                # sözcüklerini de arıyordu ve 58 bölüm işaretliyordu;
                # oysa neredeyse her 2. bölüm bir yayın tarihinden söz
                # eder. Aranan şey, malzemenin GERİ TUTULDUĞUNU söyleyen
                # cümledir — yani editörün anlamını bozmaması gereken yer.
                low = text.lower()
                if any(w in low for w in
                       ("not set out", "not reproduced", "not described",
                        "stays with", "stay with", "remain with",
                        "is not printed", "no amulet", "not drawn on",
                        "left where they are", "does not go near",
                        "without describing", "not named or located")):
                    add(f"{cid}/{key}", "yaşayan gelenek kısıtı",
                        "doğruluk doğallıktan önce gelir; anlamı "
                        "değiştirmeden düzeltin")

    for _group, key, title, _pages in MATTER_SECTIONS:
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if item.get("body"):
            add(f"matter/{key}", "yeni proza",
                f"“{title}” — kitabın en çok okunan sayfalarından")

    return flags


def render_markdown(book: dict, spec: dict) -> str:
    by_id = {c["id"]: c for c in spec["creatures"]}
    order = sorted(by_id.values(), key=lambda c: c.get("number", 0))
    classes = {c["id"]: c for c in spec["classes"]}

    L = [f"# {BOOK_TITLE}", "", f"*{BOOK_SUBTITLE}*", "",
         f"{AUTHOR} · line-editing copy", "",
         "> This file is the manuscript prepared for a native-speaker line",
         "> edit. Section headings are working labels and are not printed in",
         "> the book. Please edit in place and keep the section order.", "",
         "---", ""]

    for _group, key, title, _pages in MATTER_SECTIONS:
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if not item.get("body"):
            continue
        L += [f"## {title}", f"`matter/{key}`", "", item["body"], "", "---", ""]

    for cid, body in sorted((book.get("classOpenings") or {}).items()):
        L += [f"## Class {cid} — opening", f"`class/{cid}`", "",
              body, "", "---", ""]
    for cid, body in sorted((book.get("kinOpenings") or {}).items()):
        L += [f"## Kin family {cid} — opening", f"`kin/{cid}`", "",
              body, "", "---", ""]

    for rec in order:
        entry = (book.get("entries") or {}).get(rec["id"])
        if not entry:
            continue
        cls = classes.get(rec["class"], {})
        L += [f"## {rec['number']:03d} · {rec['name']}",
              f"`{rec['id']}` · class {rec['class']} "
              f"({cls.get('en', '')}) · {rec['tradition']}", ""]
        for key, label, _lo, _hi in ENTRY_SECTIONS:
            text = (entry.get("sections") or {}).get(key, "")
            if not text:
                continue
            L += [f"**{key}** — *{label}*", "", text, ""]
        L += ["---", ""]
    return "\n".join(L)


def render_brief(book: dict, spec: dict, flags: list[dict],
                 counts: dict) -> str:
    grouped: dict[str, list[str]] = {}
    for f in flags:
        grouped.setdefault(f["reason"], []).append(f["ref"])

    L = ["# EDITOR BRIEF — Codex Bestiarium", "",
         "> Prepared automatically by `08_BUILD/editor_pack.py`.",
         "> Do not edit this file by hand; it is regenerated.", "",
         "## What this job is", "",
         "A native-speaker line edit for **naturalness of voice**. The book",
         "was written in English by a non-native speaker, and the research,",
         "the structure and the factual content are settled. What is wanted",
         "is the ear: idiom, rhythm, and any sentence that reads as",
         "translated rather than written.", "",
         "## What is settled and should not be changed", "",
         "- **Section order.** Seven parts per entry, always in the same",
         "  order. Do not merge, reorder or add sections.",
         f"- **Word bands.** Each section has a hard band and the whole entry",
         f"  sits in {WORD_BAND[0]}–{WORD_BAND[1]} words. A rewrite that",
         "  leaves a section outside its band will fail the build. Bands are",
         "  listed against every section below.",
         "- **Facts, dates, names and citations.** Every claim is sourced.",
         "  If something looks wrong, flag it rather than fixing it.",
         "- **Diacritics.** Sīmurgh, Ḫumbaba, Àbíkú, Húli jīng. Never strip",
         "  a mark to make a word easier to set.",
         "- **Restricted material.** Where an entry says that ceremonial or",
         "  practitioner knowledge is withheld, that sentence is an ethical",
         "  commitment. Improve the English; do not make it more specific,",
         "  and do not make it shorter than the meaning requires.", "",
         "## House rules the build enforces", "",
         "These are checked mechanically on every commit. A rewrite that",
         "introduces one of them will turn the build red.", "",
         "**No hedging.** Uncertainty is handled in the source note, never",
         "in the sentence:", ""]
    L += ["- " + ", ".join(f"“{p}”" for p in FORBIDDEN_PHRASES), ""]
    L += ["**No game terminology.** This is not a game supplement:", "",
          "- " + ", ".join(f"“{p}”" for p in FORBIDDEN_GAME_TERMS), "",
          "**No unmeasurable superlatives:**", "",
          "- " + ", ".join(f"“{p}”" for p in FORBIDDEN_SUPERLATIVES), "",
          "**No softening.** The creatures are not misunderstood:", "",
          "- " + ", ".join(f"“{p}”" for p in FORBIDDEN_SOFTENERS), "",
          "**No exclamation marks anywhere in the book.**", "",
          "**Typography.** Curly quotes, en/em dashes, and `…` rather than",
          "three dots. Straight quotes fail the build.", "",
          "**Sentence rhythm.** The book aims at an average of 14–18 words a",
          "sentence. Individual sentences may be any length; the average is",
          "what is measured.", "",
          "## Sections that want human attention", "",
          "This list is derived from measurements, not from taste. It is a",
          "priority order, not a restriction: the whole manuscript is open.",
          ""]

    labels = {
        "defter dokundu": "Rewritten during the editorial pass — most "
                          "likely to read stiffly",
        "yaşayan gelenek kısıtı": "Living-tradition restriction — accuracy "
                                  "outranks fluency here",
        "bant kenarı": "Within 6% of a word band edge — compression may show",
        "ritim aykırısı": "Sentence average outside 13–19 words",
        "yeni proza": "New prose, and the most-read pages in the book",
        "diakritik yoğun": "Diacritic-dense — typographic risk",
    }
    for reason, refs in sorted(grouped.items(), key=lambda x: -len(x[1])):
        uniq = sorted(set(refs))
        L += [f"### {labels.get(reason, reason)} — {len(uniq)}", "",
              "`" + "` · `".join(uniq) + "`", ""]

    L += ["## Word bands, for reference", "", "| Section | Band |", "|---|---|"]
    for key, label, lo, hi in ENTRY_SECTIONS:
        L.append(f"| `{key}` — {label} | {lo}–{hi} |")
    L += ["", f"| **Whole entry** | **{WORD_BAND[0]}–{WORD_BAND[1]}** |", "",
          "## Scale", "",
          f"- {counts['entries']} entries · {counts['classOpenings']} class "
          f"openings · {counts['kinOpenings']} kin openings",
          f"- {counts['matter']} front/back matter sections",
          f"- {counts['words']:,} words in total".replace(",", ","),
          f"- {counts['flagged']} sections flagged below, out of "
          f"{counts['sections']}", "",
          "## How to return the work", "",
          "Edit the manuscript file in place and return it. Every change is",
          "logged into the project's edit ledger with a reason, so a short",
          "note on anything non-obvious is worth more than it looks.", ""]
    return "\n".join(L)


def gather(book: dict, spec: dict) -> tuple[list[dict], dict]:
    summary_path = os.path.join(ROOT, "01_SOURCE", "edits_summary.json")
    touched: set[str] = set()
    if os.path.exists(summary_path):
        with open(summary_path, encoding="utf-8") as fh:
            data = json.load(fh)
        # Özet madde kimliği taşır, bölüm taşımaz; bölüm defterdedir ve
        # defter depoda değildir. İki kaynak varsa ikisi de okunur.
        edits_path = os.path.join(ROOT, "01_SOURCE", "edits.json")
        if os.path.exists(edits_path):
            with open(edits_path, encoding="utf-8") as fh:
                for e in json.load(fh):
                    touched.add(f"{e['id']}/{e['section']}")
        else:
            for cid in data.get("entriesTouched", []):
                touched.add(f"{cid}/*")

    flags = flag_sections(book, spec, touched)

    sections = sum(
        1 for e in (book.get("entries") or {}).values()
        for v in (e.get("sections") or {}).values() if v)
    words = sum(
        word_count(v) for e in (book.get("entries") or {}).values()
        for v in (e.get("sections") or {}).values())
    for group in ("frontMatter", "backMatter"):
        for item in (book.get(group) or {}).values():
            words += word_count(item.get("body", "") if isinstance(item, dict)
                                else item)
    for group in ("classOpenings", "kinOpenings"):
        for body in (book.get(group) or {}).values():
            words += word_count(body)

    counts = {
        "entries": len(book.get("entries") or {}),
        "classOpenings": len(book.get("classOpenings") or {}),
        "kinOpenings": len(book.get("kinOpenings") or {}),
        "matter": len(book.get("frontMatter") or {}) +
                  len(book.get("backMatter") or {}),
        "sections": sections,
        "words": words,
        "flagged": len({f["ref"] for f in flags}),
    }
    return flags, counts


def build_docx(book: dict, spec: dict) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        print("ATLANDI: python-docx yok — DOCX üretilmedi "
              "(Markdown paketi yeterlidir).")
        return False

    by_id = {c["id"]: c for c in spec["creatures"]}
    order = sorted(by_id.values(), key=lambda c: c.get("number", 0))
    doc = Document()
    cp = doc.core_properties
    cp.title = f"{BOOK_TITLE}: {BOOK_SUBTITLE}"
    cp.author = AUTHOR
    cp.subject = "Line-editing copy"
    cp.language = "en-GB"

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    sec.left_margin = sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.5   # satır arası — editör payı

    doc.add_heading(BOOK_TITLE, level=0)
    doc.add_paragraph(BOOK_SUBTITLE)
    doc.add_paragraph(f"{AUTHOR} — line-editing copy")

    for _group, key, title, _pages in MATTER_SECTIONS:
        item = (book.get(matter_group(key)) or {}).get(key) or {}
        if not item.get("body"):
            continue
        doc.add_page_break()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"matter/{key}").italic = True
        for para in item["body"].split("\n\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("## "):
                doc.add_heading(para[3:].strip(), level=2)
            else:
                doc.add_paragraph(para)

    for kind, field in (("Class", "classOpenings"), ("Kin family",
                                                     "kinOpenings")):
        for cid, body in sorted((book.get(field) or {}).items()):
            doc.add_page_break()
            doc.add_heading(f"{kind} {cid} — opening", level=1)
            for para in body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

    for rec in order:
        entry = (book.get("entries") or {}).get(rec["id"])
        if not entry:
            continue
        doc.add_page_break()
        doc.add_heading(f"{rec['number']:03d} · {rec['name']}", level=1)
        doc.add_paragraph(f"{rec['id']} · class {rec['class']} · "
                          f"{rec['tradition']}")
        for key, label, lo, hi in ENTRY_SECTIONS:
            text = (entry.get("sections") or {}).get(key, "")
            if not text:
                continue
            h = doc.add_paragraph()
            run = h.add_run(f"{key} — {label}  [{lo}–{hi} words]")
            run.bold = True
            doc.add_paragraph(text)

    os.makedirs(MANUSCRIPT_DIR, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"yazıldı: {os.path.relpath(DOCX_PATH, ROOT)}")
    return True


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--docx", action="store_true")
    ap.add_argument("--check", action="store_true")
    args = ap.parse_args()

    book = load_book()
    if book is None or not book.get("entries"):
        print("ATLANDI: metin yok — teslim paketi yazımdan sonradır.")
        return 2
    spec = load_spec()

    flags, counts = gather(book, spec)
    measure = {
        "note": "Editör teslim paketinin ÖLÇÜSÜ. Proza içermez "
                "(karar A1/D29). Üreten: 08_BUILD/editor_pack.py",
        "counts": counts,
        "flagsByReason": {},
    }
    for f in flags:
        measure["flagsByReason"][f["reason"]] = \
            measure["flagsByReason"].get(f["reason"], 0) + 1
    measure["flagsByReason"] = dict(sorted(measure["flagsByReason"].items()))
    payload = json.dumps(measure, ensure_ascii=False, indent=2) + "\n"

    if args.check:
        if not os.path.exists(MEASURE_PATH):
            print("BAYAT: 01_SOURCE/editor_pack.json yok")
            return 1
        with open(MEASURE_PATH, encoding="utf-8") as fh:
            if fh.read() != payload:
                print("BAYAT: editor_pack.json güncel değil — "
                      "python3 08_BUILD/editor_pack.py")
                return 1
        print("TAMAM: editör paketi ölçüsü güncel.")
        return 0

    os.makedirs(MANUSCRIPT_DIR, exist_ok=True)
    with open(MD_PATH, "w", encoding="utf-8") as fh:
        fh.write(render_markdown(book, spec))
    with open(BRIEF_PATH, "w", encoding="utf-8") as fh:
        fh.write(render_brief(book, spec, flags, counts))
    with open(MEASURE_PATH, "w", encoding="utf-8") as fh:
        fh.write(payload)

    print(f"yazıldı: {os.path.relpath(MD_PATH, ROOT)}")
    print(f"yazıldı: {os.path.relpath(BRIEF_PATH, ROOT)}")
    print(f"yazıldı: {os.path.relpath(MEASURE_PATH, ROOT)}")
    if args.docx:
        build_docx(book, spec)

    print(f"\n  {counts['entries']} madde · {counts['words']:,} kelime · "
          f"{counts['flagged']}/{counts['sections']} bölüm işaretlendi"
          .replace(",", "."))
    for reason, n in measure["flagsByReason"].items():
        print(f"    {n:>4}  {reason}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
