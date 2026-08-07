# -*- coding: utf-8 -*-
"""KDP-uploadable DOCX manuscript + Kindle-ready reflowable EPUB 3."""
import os, sys, html, uuid
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import model, matter as M

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)

# ══════════════════════════════ DOCX ══════════════════════════════
def build_docx(parts, civs, order):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    cp = doc.core_properties
    cp.title, cp.author, cp.subject = f"{M.TITLE}: {M.SUBTITLE}", M.AUTHOR, "World mythology retold"
    cp.language = 'en-GB'

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(6), Inches(9)
    sec.left_margin = sec.right_margin = Inches(0.75)
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.gutter = Inches(0.125)

    def setfont(style, name, size, italic=False, bold=False, color=None):
        f = style.font
        f.name, f.size, f.italic, f.bold = name, Pt(size), italic, bold
        if color: f.color.rgb = RGBColor.from_string(color)
        rpr = style.element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), name)

    st = doc.styles
    setfont(st['Normal'], 'Garamond', 11.5)
    st['Normal'].paragraph_format.space_after = Pt(0)
    st['Normal'].paragraph_format.line_spacing = 1.18
    st['Normal'].paragraph_format.widow_control = True

    for nm, sz, it in (('Heading 1', 20, False), ('Heading 2', 15, False), ('Heading 3', 12, True)):
        s = st[nm]; setfont(s, 'Garamond', sz, italic=it, color='111111')
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(0); s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    def para(text, style=None, align=None, indent=None, italic=False,
             size=None, space_before=0, space_after=0, keep=False):
        p = doc.add_paragraph(style=style)
        r = p.add_run(text)
        if italic: r.italic = True
        if size: r.font.size = Pt(size)
        pf = p.paragraph_format
        if align is not None: pf.alignment = align
        if indent is not None: pf.first_line_indent = Inches(indent)
        pf.space_before, pf.space_after = Pt(space_before), Pt(space_after)
        pf.keep_with_next = keep
        return p

    def pagebreak():
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    C, J = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY

    # front matter
    para(M.TITLE, align=C, size=26, space_before=120); pagebreak()
    para(M.TITLE, align=C, size=26, space_before=90)
    para(M.SUBTITLE, align=C, size=13, italic=True, space_before=8)
    para(M.AUTHOR, align=C, size=14, space_before=90)
    para(M.IMPRINT, align=C, size=11, space_before=60); pagebreak()
    for _, t in M.COPYRIGHT: para(t, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, space_after=6)
    pagebreak()
    para(M.DEDICATION[0][1], align=C, italic=True, size=12, space_before=140); pagebreak()

    para('Contents', style='Heading 1'); para('', space_after=6)
    for p in parts:
        if p['kind'] == 'civ':
            para(p['full'], align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, space_before=8, space_after=2)
        else:
            para('    ' + p['title'], align=WD_ALIGN_PARAGRAPH.LEFT, size=10, space_after=1)
    pagebreak()

    for _, title, plist in model.front_sections():
        para(title, style='Heading 1'); para('', space_after=8)
        for i, t in enumerate(plist):
            para(t, align=J, indent=0 if i == 0 else 0.22)
        pagebreak()

    # body
    for p in parts:
        if p['kind'] == 'civ':
            para('', space_before=120)
            para(p['full'], style='Heading 1')
            para(f"{p['name']} · {p['epoch']}", align=C, italic=True, size=11, space_after=10)
            para(p['desc'], align=C, size=11, space_after=6)
            pagebreak()
        else:
            para(p['title'], style='Heading 2')
            para(p['subtitle'], style='Heading 3'); para('', space_after=8)
            first = True
            for kind, text in p['blocks']:
                if kind == 'p':
                    para(text, align=J, indent=0 if first else 0.22); first = False
                elif kind == 'q':
                    q = para(text, align=J, italic=True, size=11, space_before=8, space_after=8)
                    q.paragraph_format.left_indent = Inches(0.35)
                    q.paragraph_format.right_indent = Inches(0.35)
                else:
                    para('*   *   *', align=C, space_before=10, space_after=10)
            pagebreak()

    for tag, title, plist in model.back_sections(civs, order, parts):
        para(title, style='Heading 1'); para('', space_after=8)
        for i, t in enumerate(plist):
            if tag == 'Reference':
                para(t, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, space_after=3)
            else:
                para(t, align=J, indent=0 if i == 0 else 0.22)
        pagebreak()

    out = os.path.join(ROOT, 'CODEX_MYTHOLOGICA_MANUSCRIPT.docx')
    doc.save(out)
    return out

# ══════════════════════════════ EPUB ══════════════════════════════
CSS = """@charset "utf-8";
body { margin:0 5%; line-height:1.5; text-align:justify; font-family:Georgia,'Iowan Old Style',serif;
       widows:2; orphans:2; hyphens:auto; -epub-hyphens:auto; }
h1.book  { text-align:center; font-size:1.9em; margin:22% 0 0.3em; font-weight:normal; letter-spacing:.06em; }
p.sub    { text-align:center; font-style:italic; font-size:1.05em; color:#444; margin:0 0 3em; }
p.author { text-align:center; font-size:1.1em; letter-spacing:.08em; margin:2.5em 0 0; }
p.imprint{ text-align:center; font-size:.85em; color:#555; margin:3em 0 0; }
p.rights { font-size:.82em; text-align:left; margin:0 0 .8em; color:#333; }
p.ded    { text-align:center; font-style:italic; margin:35% 0 0; color:#333; }
h1.civ   { text-align:center; font-size:1.6em; font-weight:normal; letter-spacing:.09em;
           margin:22% 0 .2em; page-break-before:always; }
p.civmeta{ text-align:center; font-style:italic; color:#555; margin:0 0 1.4em; font-size:.95em; }
p.civdesc{ text-align:center; color:#333; margin:0 6% 0; }
h1.story { text-align:center; font-size:1.35em; font-weight:normal; margin:8% 0 .15em;
           letter-spacing:.04em; page-break-before:always; }
p.stsub  { text-align:center; font-style:italic; color:#555; margin:0 0 1.8em; font-size:.95em; }
h1.fm    { text-align:center; font-size:1.4em; font-weight:normal; margin:10% 0 1.4em;
           letter-spacing:.05em; page-break-before:always; }
p        { margin:0; text-indent:1.2em; }
p.first  { text-indent:0; }
p.first::first-line { font-variant:small-caps; letter-spacing:.03em; }
blockquote { margin:1.1em 2em; font-style:italic; text-align:justify; }
blockquote p { text-indent:0; }
p.aster  { text-align:center; text-indent:0; margin:1.3em 0; color:#666; letter-spacing:.5em; }
nav ol   { list-style:none; padding-left:0; }
nav ol ol{ padding-left:1.2em; }
ul.civs  { list-style:none; padding-left:0; }
ul.civs li { margin:0 0 .45em; text-indent:-1em; padding-left:1em; }
"""

def x(t): return html.escape(t, quote=False)

def build_epub(parts, civs, order):
    from ebooklib import epub
    bk = epub.EpubBook()
    bk.set_identifier(str(uuid.uuid5(uuid.NAMESPACE_URL, 'codex-mythologica-2026')))
    bk.set_title(f"{M.TITLE}: {M.SUBTITLE}")
    bk.set_language('en')
    bk.add_author(M.AUTHOR)
    bk.add_metadata('DC', 'publisher', M.IMPRINT)
    bk.add_metadata('DC', 'date', M.YEAR)
    bk.add_metadata('DC', 'description',
                    'Seventy-six myths from nineteen civilisations, retold in full — Greek, Norse, '
                    'Egyptian, Hindu, Japanese, Celtic, Mesopotamian, Aztec, Maya, Chinese, Korean, '
                    'Slavic, West African, Persian, Polynesian, Inuit, Turkic, Roman and Arabian.')
    for s in ('Mythology', 'Folklore', 'World Literature', 'Short Stories', 'Legends'):
        bk.add_metadata('DC', 'subject', s)

    css = epub.EpubItem(uid='style', file_name='style/book.css',
                        media_type='text/css', content=CSS)
    bk.add_item(css)

    def page(uid, name, title, body):
        # ebooklib supplies the XHTML shell; hand it a body fragment only
        c = epub.EpubHtml(uid=uid, file_name=name, title=title, lang='en')
        c.set_content(body)
        c.add_item(css); bk.add_item(c); return c

    spine, toc_top = [], []

    spine.append(page('title', 'title.xhtml', M.TITLE,
        f'<h1 class="book">{x(M.TITLE)}</h1><p class="sub">{x(M.SUBTITLE)}</p>'
        f'<p class="author">{x(M.AUTHOR)}</p><p class="imprint">{x(M.IMPRINT)}</p>'))
    spine.append(page('rights', 'copyright.xhtml', 'Copyright',
        ''.join(f'<p class="rights">{x(t)}</p>' for _, t in M.COPYRIGHT)))
    spine.append(page('ded', 'dedication.xhtml', 'Dedication',
        f'<p class="ded">{x(M.DEDICATION[0][1])}</p>'))

    for key, title, plist in model.front_sections():
        c = page(key.lower(), f'{key.lower()}.xhtml', title,
                 f'<h1 class="fm">{x(title)}</h1>' +
                 ''.join(f'<p class="{"first" if i==0 else ""}">{x(t)}</p>' for i, t in enumerate(plist)))
        spine.append(c); toc_top.append(c)

    civ_sections = []
    cur_children, cur_civ = [], None
    for p in parts:
        if p['kind'] == 'civ':
            if cur_civ: civ_sections.append((cur_civ, cur_children)); cur_children = []
            cur_civ = page(f"civ-{p['id']}", f"civ-{p['id']}.xhtml", p['full'],
                f'<h1 class="civ">{x(p["full"])}</h1>'
                f'<p class="civmeta">{x(p["name"])} &#183; {x(p["epoch"])}</p>'
                f'<p class="civdesc">{x(p["desc"])}</p>')
            spine.append(cur_civ)
        else:
            body = [f'<h1 class="story">{x(p["title"])}</h1>',
                    f'<p class="stsub">{x(p["subtitle"])}</p>']
            first = True
            for kind, text in p['blocks']:
                if kind == 'p':
                    body.append(f'<p class="{"first" if first else ""}">{x(text)}</p>'); first = False
                elif kind == 'q':
                    body.append(f'<blockquote><p>{x(text)}</p></blockquote>')
                else:
                    body.append('<p class="aster">* * *</p>')
            c = page(p['id'], f"s-{p['id']}.xhtml", p['title'], ''.join(body))
            spine.append(c); cur_children.append(c)
    if cur_civ: civ_sections.append((cur_civ, cur_children))

    for tag, title, plist in model.back_sections(civs, order, parts):
        if tag == 'Reference':
            body = f'<h1 class="fm">{x(title)}</h1><ul class="civs">' + \
                   ''.join(f'<li>{x(t)}</li>' for t in plist) + '</ul>'
        else:
            body = f'<h1 class="fm">{x(title)}</h1>' + \
                   ''.join(f'<p class="{"first" if i==0 else ""}">{x(t)}</p>' for i, t in enumerate(plist))
        c = page(tag.lower(), f'{tag.lower()}.xhtml', title, body)
        spine.append(c); toc_top.append(c)

    bk.toc = tuple(toc_top[:2] +
                   [(epub.Section(sec.title), tuple(kids)) for sec, kids in civ_sections] +
                   toc_top[2:])
    bk.add_item(epub.EpubNcx()); bk.add_item(epub.EpubNav())
    bk.spine = ['nav'] + spine
    out = os.path.join(ROOT, 'CODEX_MYTHOLOGICA.epub')
    epub.write_epub(out, bk, {'epub3_landmark': False, 'epub3_pages': False})
    return out

if __name__ == '__main__':
    B, civs, order, parts = model.load()
    d = build_docx(parts, civs, order); print('DOCX:', d, os.path.getsize(d) // 1024, 'KB')
    e = build_epub(parts, civs, order); print('EPUB:', e, os.path.getsize(e) // 1024, 'KB')
